# This is a sample Python script.
from docx import Document
# Press ⌃R to execute it or replace it with your code.
# Press Double ⇧ to search everywhere for classes, files, tool windows, actions, and settings.


def print_hi(name):
    document = Document('data/INCS 101 “Programming I”.docx')
    for p in document.paragraphs:
        print(p.text)
    for table in document.tables:
        for row in range(len(table.rows)):
            for col in range(len(table.columns)):
                print(table.cell(row, col).text)



    #document.save('new-file-name.docx')

# Press the green button in the gutter to run the script.
if __name__ == '__main__':
    print_hi('PyCharm')

# See PyCharm help at https://www.jetbrains.com/help/pycharm/
