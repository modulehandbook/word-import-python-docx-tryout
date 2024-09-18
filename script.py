import json
from docx import Document
import re
import sys


def extract_info(document):
    document = Document(document)
    paragraph_data = extract_info_from_paragraphs(document)
    table_data = extract_info_from_tables(document)

    data = {**paragraph_data, **table_data}

    return data


def extract_info_from_paragraphs(document):
    extracted_info = {}
    collecting = False
    current_field = None
    collecting_text = ""

    p_keywords = {
        "Course Aim": "mission",
        "Knowledge and Understanding": "skills_knowledge_understanding",
        "Intellectual Skills": "skills_intellectual",
        "Professional and Practical Skills": "skills_practical",
        "General and Transferable Skills": "skills_general",
        "General and Transferrable Skills": "skills_general", #typo in some files
        "Learning and Teaching Methods": "methods",
        "Facilities Required For Teaching and Learning": "equipment",
        "References": "literature",
        "Head of the Department:": "department_head",
        "Course Coordinator:": "responsible_person",
        "E-mail:": "mail"
    }
    non_p_keywords = [
        "Course Content",
        "Assessment",
        "Learning Outcomes",
        "C- Administrative Information"
    ]

    all_keywords = set(p_keywords.keys()) | set(non_p_keywords)
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]

    extracted_info["name"] = re.findall(r'[“"](.*?)[”"]', paragraphs[0])[0]
    extracted_info["code"] = re.findall(r'^[^"“]*', paragraphs[0])[0].strip()

    i = 0
    while i < len(paragraphs):
        para_text = paragraphs[i]

        if collecting:
            collecting_text += "\n" + para_text
            # stop collecting if next paragraph is keyword
            if i + 1 < len(paragraphs):
                next_para_text = paragraphs[i + 1]
                if any(keyword in next_para_text for keyword in all_keywords):
                    extracted_info[current_field] = collecting_text.strip()
                    collecting = False
                    current_field = None
                    collecting_text = ""
        else:
            # check if current paragraph contains p_keywords
            for keyword, field in p_keywords.items():
                if keyword.lower() in para_text.lower():
                    if ":" in para_text:
                        extracted_info[field] = para_text.split(":", 1)[1].strip()
                        break
                    collecting = True
                    current_field = field
                    collecting_text = ""
                    break
        i += 1

    #case where last collected text needs to be added
    if collecting and current_field:
        extracted_info[current_field] = collecting_text.strip()

    return extracted_info


def extract_info_from_tables(document):
    table_general = document.tables[0]
    extracted_info = {"semester": table_general.rows[0].cells[1].text,
                      "required": table_general.rows[5].cells[1].text,
                      "ects": table_general.rows[8].cells[1].text,
                      "lectureHrs": 0,
                      "tutorialHrs": 0,
                      "labHrs": 0,
                      "contents": contents_table_to_json(document.tables[1]),
                      "examination": assessment_table_to_json(document.tables[2])
                      }

    temp = re.split(r' \+ |, ', table_general.rows[6].cells[1].text)
    for session in temp:
        if 'lecture' in session.lower():
            extracted_info["lectureHrs"] += calculate_sws(session)
        elif 'tutorial' in session.lower():
            extracted_info["tutorialHrs"] += calculate_sws(session)
        elif 'lab' in session.lower():
            extracted_info["labHrs"] += calculate_sws(session)

    return extracted_info


def contents_table_to_json(table):
    table_data = []

    headers = [cell.text.strip() for cell in table.rows[0].cells]
    sub_headers = [cell.text.strip() for cell in table.rows[1].cells]

    for row in table.rows[2:]:
        row_data = {}
        for index, cell in enumerate(row.cells):
            header = headers[index]
            if index < len(sub_headers) and sub_headers[index]:  # if sub-header in column
                if header not in row_data:
                    row_data[header] = {}
                row_data[header][sub_headers[index]] = cell.text.strip()
            else:
                row_data[header] = cell.text.strip()

        table_data.append(row_data)
    return table_data


def assessment_table_to_json(table):
    extracted_info = []
    headers = [cell.text.strip() for cell in table.rows[0].cells]

    for row in table.rows[1:]:
        row_data = {}
        for index, cell in enumerate(row.cells):
            row_data[headers[index]] = cell.text.strip()
        extracted_info.append(row_data)

    return extracted_info


def calculate_sws(session_string):
    if re.search(r'bi\s*-?\s*weekly', session_string.lower()):
        number = int(session_string.split()[0])
        return number
    else:
        number = int(session_string.split()[0])
        return number * 2

#if __name__ == '__main__':
#    document_path = 'data/INCS 101 “Programming I”.docx'
#    info = extract_info(document_path)
#    with open('combined_data.json', 'w') as json_file:
#        json.dump(info, json_file, indent=4)

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python script_name.py <document_path>")
        sys.exit(1)

    document_path = sys.argv[1]
    info = extract_info(document_path)
    print(json.dumps(info, indent=4))
